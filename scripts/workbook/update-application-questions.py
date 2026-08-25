#!/usr/bin/env python3

import argparse
import json
import re
from pathlib import Path

from docx import Document


def replace_questions(source: Path, destination: Path, questions: list[str]) -> None:
    document = Document(source)

    if source.name == "lesson-07-handout.docx":
        for paragraph in document.paragraphs:
            if paragraph.text.strip() == "Romans 6:1-2; Romans 6:1-14":
                paragraph.runs[0].text = "Romans 6:1-14"
                for extra_run in paragraph.runs[1:]:
                    extra_run.text = ""
                break
            if paragraph.text.strip() == "Romans 6:1-14":
                break
        else:
            raise ValueError(f"{source.name}: could not find the Romans 6 reference")

    question_paragraphs = [
        paragraph
        for paragraph in document.paragraphs
        if re.match(r"^[1-4]\.\s+", paragraph.text)
    ]

    if len(question_paragraphs) != 4:
        raise ValueError(f"{source.name}: expected 4 question paragraphs, found {len(question_paragraphs)}")

    for number, (paragraph, question) in enumerate(zip(question_paragraphs, questions), start=1):
        if len(paragraph.runs) < 2:
            raise ValueError(f"{source.name}: question {number} does not preserve the expected number/text runs")
        paragraph.runs[0].text = f"{number}.  "
        paragraph.runs[1].text = question
        for extra_run in paragraph.runs[2:]:
            extra_run.text = ""

    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


def main() -> None:
    parser = argparse.ArgumentParser(description="Replace the four application questions in each Growing in Grace handout.")
    parser.add_argument("--source-dir", type=Path, required=True)
    parser.add_argument("--output-dir", type=Path, required=True)
    parser.add_argument(
        "--questions",
        type=Path,
        default=Path("content/growing-in-grace-application-questions.json"),
    )
    args = parser.parse_args()

    question_sets = json.loads(args.questions.read_text())
    if sorted(map(int, question_sets)) != list(range(1, 29)):
        raise ValueError("Question source must contain lessons 1 through 28")

    for lesson_number in range(1, 29):
        questions = question_sets[str(lesson_number)]
        if len(questions) != 4 or any(not question.strip() for question in questions):
            raise ValueError(f"Lesson {lesson_number}: expected exactly four non-empty questions")

        filename = f"lesson-{lesson_number:02d}-handout.docx"
        replace_questions(
            args.source_dir / filename,
            args.output_dir / filename,
            questions,
        )
        print(f"Updated {filename}")


if __name__ == "__main__":
    main()
